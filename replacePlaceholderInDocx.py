from docx import Document
import csv
import sys
import os
from datetime import datetime
from pathlib import Path

#print('Number of arguments:', len(sys.argv), 'arguments.')
#print('Argument List:', str(sys.argv))

#pathToMappingFile = "./ronny.csv"
pathToMappingFile = "/data/mapping.csv"
mydict = {}

if not Path(pathToMappingFile).is_file():
    print("File 'mapping.csv' does not exist in mounted directory. Exiting!")
    print("Directory is: " + pathToMappingFile)
    exit()

#with open(pathToMappingFile, mode='r', encoding= 'unicode_escape') as infile:
with open(pathToMappingFile, mode='r') as infile:
    reader = csv.reader(infile,delimiter=";")
    #for rows in reader:
        #print(rows)
    mydict = {rows[0]:rows[1] for rows in reader}

mydict['<DATE>'] = datetime.today().strftime('%d.%m.%Y')
#print(mydict['<DATE>'])

document = Document("NasuniEvaluationTestPlan.docx")
if os.environ['POC_DOC_TYPE'] == 'naa':
    document = Document("NAAEvalTestPlan.docx")

#print(mydict)

print(document)

for paragraph in document.paragraphs:
    for run in paragraph.runs:
        #print("-------")
        #print(run.text)
        for placeholder in mydict:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder,mydict[placeholder])
                #print("New text: " + run.text)

for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for placeholder in mydict:
                if placeholder in cell.text:
                    #print("-------")
                    #print("vorher:" + cell.text)
                    cell.text = cell.text.replace(placeholder,mydict[placeholder])
                    #print("nachher:" + cell.text)

#print(document.paragraphs[8].runs[0].text)


#Dictionary = {'sea': 'ocean'}
filename = "/data/Nasuni Evaluation Test Plan - " + mydict.get("<CUSTOMER NAME>") + " - " + datetime.today().strftime('%Y-%m-%d') + ".docx"
#filename = "./Nasuni Evaluation Test Plan - " + mydict.get("<CUSTOMER NAME>") + " - " + datetime.today().strftime('%Y-%m-%d') + ".docx"
document.save(filename)

